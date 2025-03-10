import os
import sys
import shutil
import yaml
import logging
import requests
from loguru import logger
from datetime import datetime

def setup_logging():
    """Configure logging for the application."""
    # Create logs directory if it doesn't exist
    os.makedirs("logs", exist_ok=True)
    
    # Configure loguru
    log_format = "<green>{time:YYYY-MM-DD HH:mm:ss}</green> | <level>{level: <8}</level> | <cyan>{name}</cyan>:<cyan>{function}</cyan>:<cyan>{line}</cyan> - <level>{message}</level>"
    
    # Remove default handler
    logger.remove()
    
    # Add console handler
    logger.add(sys.stderr, format=log_format, level="INFO")
    
    # Add file handler with rotation
    log_file = f"logs/job_application_{datetime.now().strftime('%Y%m%d')}.log"
    logger.add(log_file, format=log_format, level="DEBUG", rotation="10 MB", retention="1 week")
    
    logger.info("Logging configured")

def create_directory_structure():
    """Create the necessary directory structure for the application."""
    directories = [
        "config",
        "data",
        "logs",
        "src"
    ]
    
    for directory in directories:
        os.makedirs(directory, exist_ok=True)
    
    logger.info("Directory structure created")

def create_sample_resume():
    """
    Set up resume for the application.
    
    This function checks for a PDF resume in the data directory.
    If not found, it creates a basic Word document template as fallback.
    """
    # Check for PDF resume
    pdf_resume_path = "data/resume.pdf"
    docx_resume_path = "data/resume_template.docx"
    
    # If PDF resume exists, make a backup
    if os.path.exists(pdf_resume_path):
        logger.info(f"Found existing PDF resume at {pdf_resume_path}")
        
        # Create a backup of the original PDF resume if it doesn't exist
        pdf_backup_path = "data/resume_original.pdf"
        if not os.path.exists(pdf_backup_path):
            shutil.copy2(pdf_resume_path, pdf_backup_path)
            logger.info(f"Created backup of original PDF resume at {pdf_backup_path}")
        
        return
    
    # If PDF doesn't exist, create a basic Word document as fallback
    if not os.path.exists(docx_resume_path):
        _create_basic_word_template(docx_resume_path)
        logger.warning("No PDF resume found. Please add your PDF resume to data/resume.pdf")
    else:
        logger.info(f"Found existing Word resume template at {docx_resume_path}")

def _create_basic_word_template(output_path):
    """Create a basic Word document resume template as fallback."""
    try:
        from docx import Document
        
        doc = Document()
        
        # Add name and contact info
        doc.add_paragraph("YOUR NAME").bold = True
        contact_info = doc.add_paragraph()
        contact_info.add_run("Email: your.email@example.com | Phone: (123) 456-7890 | LinkedIn: linkedin.com/in/yourprofile")
        
        # Add professional summary
        doc.add_paragraph("Professional Summary").bold = True
        doc.add_paragraph("Experienced professional with a background in software development and a passion for technology.")
        
        # Add skills section
        doc.add_paragraph("Skills").bold = True
        skills = ["Python", "JavaScript", "React", "Node.js", "SQL", "AWS", "Docker"]
        for skill in skills:
            doc.add_paragraph(skill, style='ListBullet')
        
        # Add experience section
        doc.add_paragraph("Experience").bold = True
        exp = doc.add_paragraph()
        exp.add_run("Software Engineer, Example Company").bold = True
        exp.add_run("\nJan 2020 - Present")
        doc.add_paragraph("Developed and maintained web applications using modern technologies.")
        
        # Add education section
        doc.add_paragraph("Education").bold = True
        edu = doc.add_paragraph()
        edu.add_run("Bachelor of Science in Computer Science").bold = True
        edu.add_run("\nUniversity Name, 2019")
        
        # Save the document
        doc.save(output_path)
        logger.info(f"Created basic Word resume template at {output_path}")
        return True
    except Exception as e:
        logger.error(f"Error creating Word resume template: {str(e)}")
        return False

def get_resume_path():
    """
    Get the path to the resume file.
    
    Returns:
        str: Path to the resume file (PDF preferred, Word as fallback)
    """
    pdf_path = "data/resume.pdf"
    docx_path = "data/resume_template.docx"
    
    if os.path.exists(pdf_path):
        return pdf_path
    elif os.path.exists(docx_path):
        return docx_path
    else:
        logger.error("No resume file found")
        return None 

def get_cover_letter_path():
    """Get the path to the base cover letter."""
    # Check for PDF cover letter
    pdf_path = "data/base_cover_letter.pdf"
    if os.path.exists(pdf_path):
        return pdf_path
    
    # Check for text cover letter
    txt_path = "data/base_cover_letter.txt"
    if os.path.exists(txt_path):
        return txt_path
    
    return None 

def check_ollama_available(config):
    """Check if Ollama is available and working."""
    import requests
    
    # Get Ollama host from environment or use default
    ollama_host = os.environ.get('OLLAMA_HOST', 'http://localhost:11434')
    ollama_model = config.get('llm', {}).get('ollama_model', 'llama2')
    
    try:
        # Check if Ollama server is running
        url = f"{ollama_host}/api/tags"
        response = requests.get(url)
        
        if response.status_code == 200:
            # Check if the model exists
            models = response.json().get('models', [])
            if any(model.get('name') == ollama_model for model in models):
                logger.info(f"Ollama is available with model {ollama_model}")
                return True
            else:
                logger.warning(f"Model {ollama_model} not found in Ollama. Please run: ollama pull {ollama_model}")
                return False
        else:
            logger.error(f"Ollama server returned status code: {response.status_code}")
            return False
    except Exception as e:
        logger.error(f"Error connecting to Ollama: {str(e)}")
        logger.info("Make sure Ollama is installed and running")
        return False

def debug_env_vars():
    """Debug environment variables related to OpenAI."""
    import os
    
    # Check Ollama host
    ollama_host = os.environ.get('OLLAMA_HOST')
    if ollama_host:
        logger.debug(f"OLLAMA_HOST found in environment: {ollama_host}")
    else:
        logger.warning("OLLAMA_HOST not found in environment variables, using default: http://localhost:11434")
    
    # Check if python-dotenv is installed and try to load .env file manually
    try:
        from dotenv import load_dotenv
        load_dotenv()
        
        # Check if Ollama host is available after loading .env
        ollama_host_after = os.environ.get('OLLAMA_HOST')
        if ollama_host_after and not ollama_host:
            logger.info(f"OLLAMA_HOST loaded from .env file: {ollama_host_after}")
        elif not ollama_host_after:
            logger.warning("OLLAMA_HOST not found even after loading .env file, using default")
    except ImportError:
        logger.warning("python-dotenv not installed, can't load .env file automatically") 

def prompt_yes_no(question, default="yes"):
    """Ask a yes/no question via input() and return the answer as a boolean.
    
    Args:
        question (str): The question to ask
        default (str): The default answer if the user just hits Enter
        
    Returns:
        bool: True for "yes" or False for "no"
    """
    valid = {"yes": True, "y": True, "ye": True, "no": False, "n": False}
    if default is None:
        prompt = " [y/n] "
    elif default == "yes":
        prompt = " [Y/n] "
    elif default == "no":
        prompt = " [y/N] "
    else:
        raise ValueError(f"Invalid default answer: '{default}'")
    
    while True:
        sys.stdout.write(question + prompt)
        choice = input().lower()
        if choice == '':
            return valid[default]
        elif choice in valid:
            return valid[choice]
        else:
            sys.stdout.write("Please respond with 'yes' or 'no' (or 'y' or 'n').\n") 